import os
import docx
import pandas as pd
import shutil
path='C:\\Users\\ghosh\\.PyCharmEdu2018.2\\config\\scratches\\'
print(os.getcwd())
# os.mkdir('c:\\users\\ghosh\\Desktop\\Text Mining and NLP',755)
cwdpath='c:\\users\\ghosh\\Desktop\\Text Mining and NLP\\'
os.chdir(cwdpath)
print(os.getcwd())

os.chdir(path)

if os.path.exists(path+"Welcome.txt"):
  os.remove("Welcome.txt")
else:
  print("The file does not exist")

with open(path+'greetings.txt','w') as fp:
    fp.write('Welcome toText Mining and Natural Language Processing')

os.rename('greetings.txt','Welcome.txt')

if os.path.exists(cwdpath+"Welcome.txt"):
  os.remove(cwdpath+"Welcome.txt")
else:
  print("The file does not exist")

shutil.move(path+'Welcome.txt',cwdpath)


docpath='nltk_datasets\\47y9kaz16i\\698_m1_datasets_v1.0\\NLP.docx'
csvpath='nltk_datasets\\47y9kaz16i\\698_m1_datasets_v1.0\\EmployeeDetails.csv'

docu1=docx.Document(path+docpath)

all_paras = docu1.paragraphs
print(len(all_paras))

for i in all_paras:
    print(i.text)
    print(len(str(i.text).split(' ')))

append_content=open(cwdpath+"Welcome.txt",'r').read()
print(append_content,'\n')

para=docu1.paragraphs[0]
para.add_run(append_content)
docu1.save(path+docpath)

csv_df=pd.read_csv(path+csvpath,sep=',',header='infer')
print(csv_df)


# new_df=csv_df.iloc[:,0]
# csv_df['new_col']=new_df.apply(lambda x :(str(x).split(' ')))
# print('\n\n and now:\n',new_df)


csv_df['First_name']=csv_df['Name '].str.split(' ',expand=True).iloc[:,0]
csv_df['Last_name']=csv_df['Name '].str.split(' ',expand=True).iloc[:,1]


csv_df.drop(['Name '],inplace=True,axis=1)
csv_df["Salary"]=csv_df["Salary"]*1.1
print(csv_df)

csv_df.to_csv(path+'Employee_Data.csv')

# another way od doing
# csv_df=pd.read_csv(path+csvpath,sep=',',header='infer')
# print(csv_df)
# new_df=csv_df.iloc[:,0]
# csv_df['new_col']=new_df.apply(lambda x :(str(x).split(' ')))
# print('\n\n and now:\n',new_df)
# ser_to_df=csv_df.new_col.apply(pd.Series).merge(csv_df, left_index = True, right_index = True)
# ser_to_df=ser_to_df.dropna(how='any',axis=1)
# ser_to_df=ser_to_df.drop(['new_col','Name '],axis=1)
# print(ser_to_df.columns)
# print(ser_to_df)
